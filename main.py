# from fastapi import FastAPI, File, UploadFile, Form, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# import pdfplumber
# import docx
# import requests
# import json
# import io
# import os

# app = FastAPI()

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
#     allow_headers=["*"],
# )

# def extract_text_from_pdf(file_stream):
#     text = ""
#     try:
#         file_stream.seek(0)
#         with pdfplumber.open(file_stream) as pdf:
#             for page in pdf.pages:
#                 content = page.extract_text()
#                 if content:
#                     text += content + "\n"
#     except Exception as e:
#         raise ValueError(f"Error processing PDF with pdfplumber: {e}")
#     return text

# def extract_text_from_docx(file_stream):
#     try:
#         file_stream.seek(0)
#         doc = docx.Document(file_stream)
#         return "\n".join([p.text for p in doc.paragraphs])
#     except Exception as e:
#         raise ValueError(f"Error processing DOCX with python-docx: {e}")

# async def extract_resume_text(file: UploadFile):
#     try:
#         file_content = await file.read()
#         file_stream = io.BytesIO(file_content)

#         base_filename = os.path.basename(file.filename)
#         debug_output_path = f"debug_uploaded_{base_filename}"
#         with open(debug_output_path, "wb") as f:
#             f.write(file_content)

#         if file.content_type == "application/pdf":
#             from PyPDF2 import PdfReader
#             try:
#                 reader = PdfReader(file_stream)
#                 if not reader.pages:
#                     raise ValueError("Uploaded PDF is empty or malformed.")
#                 file_stream.seek(0)
#                 return extract_text_from_pdf(file_stream)
#             except Exception as e:
#                 raise ValueError(f"Invalid PDF file: {e}")

#         elif file.content_type in [
#             "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             "application/msword"
#         ]:
#             return extract_text_from_docx(file_stream)

#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX.")

#     except HTTPException as e:
#         raise e
#     except ValueError as e:
#         raise HTTPException(status_code=422, detail=f"Error reading resume content: {e}")
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"An unexpected error occurred during resume text extraction: {str(e)}")

# @app.post("/analyze")
# async def analyze(
#     resume: UploadFile = File(...),
#     job_description: str = Form(...),
#     interview_level: str = Form("General Match (Overall Fit)")
# ):
#     try:
#         resume_text = await extract_resume_text(resume)
#     except HTTPException as e:
#         raise e
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to extract resume text: {str(e)}")

#     prompt = f"""
# You are an expert career advisor and technical interviewer. Your task is to compare a candidate's resume against a job description, focusing your analysis based on the specified interview context.

# **Interview Context:** {interview_level}

# **Resume:**
# {resume_text}

# **Job Description:**
# {job_description}

# Please analyze the resume and job description and respond strictly in JSON format. The JSON should contain:
# - `matched_skills`
# - `missing_skills`
# - `proficiency_levels`
# """

#     try:
#         ollama_response = requests.post(
#             "http://localhost:11434/api/chat",
#             json={
#                 "model": "openchat",
#                 "messages": [{"role": "user", "content": prompt}],
#                 "stream": False,
#                 "options": {
#                     "temperature": 0.3
#                 }
#             },
#             timeout=120
#         )
#         ollama_response.raise_for_status()

#         response_data = ollama_response.json()
#         content = response_data.get("message", {}).get("content", "").strip()

#         try:
#             return json.loads(content)
#         except json.JSONDecodeError:
#             return {"raw_response": content}

#     except requests.exceptions.ConnectionError:
#         raise HTTPException(status_code=503, detail="Could not connect to Ollama server at http://localhost:11434. Please ensure Ollama is running.")
#     except requests.exceptions.HTTPError as e:
#         raise HTTPException(status_code=e.response.status_code, detail=f"Error from Ollama API: {e.response.text}")
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"An unexpected error occurred during LLM analysis: {str(e)}")

import io
import json
import logging
import os
import re

import docx
import pdfplumber
import requests
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from PyPDF2 import PdfReader # Import PyPDF2 here as it's used for validation

# Initialize FastAPI app
app = FastAPI(
    title="Resume Analyzer Backend",
    description="API for analyzing resumes and generating interview preparation tips using Ollama.",
    version="1.0.0",
)

# Configure CORS to allow requests from your Gradio UI.
# In a production environment, replace "*" with the specific origin of your Gradio app (e.g., "http://localhost:7860").
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(level=logging.INFO) # Set to logging.DEBUG for more verbose output during development
logger = logging.getLogger(__name__)

# --- Helper Functions for Text Extraction ---

def extract_text_from_pdf(file_stream: io.BytesIO) -> str:
    """
    Extracts text content from a PDF file stream.

    Args:
        file_stream (io.BytesIO): The file stream of the PDF.

    Returns:
        str: The extracted text content.

    Raises:
        ValueError: If an error occurs during PDF processing.
    """
    text = ""
    try:
        file_stream.seek(0)  # Ensure stream is at the beginning
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
    except Exception as e:
        logger.error(f"Error processing PDF with pdfplumber: {e}", exc_info=True)
        raise ValueError(f"Error processing PDF with pdfplumber: {e}")
    return text

def extract_text_from_docx(file_stream: io.BytesIO) -> str:
    """
    Extracts text content from a DOCX file stream.

    Args:
        file_stream (io.BytesIO): The file stream of the DOCX.

    Returns:
        str: The extracted text content.

    Raises:
        ValueError: If an error occurs during DOCX processing.
    """
    try:
        file_stream.seek(0)  # Ensure stream is at the beginning
        doc = docx.Document(file_stream)
        # Ensure paragraphs are strings and strip whitespace, then join
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error processing DOCX with python-docx: {e}", exc_info=True)
        raise ValueError(f"Error processing DOCX with python-docx: {e}")

async def extract_resume_text(file: UploadFile) -> str:
    """
    Extracts text from an uploaded resume file (PDF or DOCX).

    Args:
        file (UploadFile): The uploaded resume file.

    Returns:
        str: The extracted text content.

    Raises:
        HTTPException: If the file type is unsupported or extraction fails.
    """
    try:
        file_content = await file.read()
        file_stream = io.BytesIO(file_content)

        # Optional: For debugging, save the uploaded file
        # base_filename = os.path.basename(file.filename)
        # with open(f"debug_uploaded_{base_filename}", "wb") as f:
        #     f.write(file_content)

        if file.content_type == "application/pdf":
            try:
                # Use PyPDF2 for initial validation (e.g., empty/malformed PDF)
                reader = PdfReader(file_stream)
                if not reader.pages:
                    raise ValueError("Uploaded PDF is empty or malformed.")
                file_stream.seek(0) # Reset stream position after PyPDF2 read for pdfplumber
                return extract_text_from_pdf(file_stream)
            except Exception as e:
                logger.error(f"Invalid PDF file during PyPDF2 validation or pdfplumber extraction: {e}", exc_info=True)
                raise ValueError(f"Invalid PDF file: {e}")

        elif file.content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            return extract_text_from_docx(file_stream)

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX.")

    except HTTPException as e:
        logger.error(f"HTTPException in extract_resume_text: {e.detail}", exc_info=True)
        raise e
    except ValueError as e:
        logger.error(f"ValueError in extract_resume_text: {e}", exc_info=True)
        raise HTTPException(status_code=422, detail=f"Error reading resume content: {e}")
    except Exception as e:
        logger.error(f"An unexpected error occurred in extract_resume_text: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during resume text extraction: {str(e)}")

# --- API Endpoints ---

@app.post("/analyze")
async def analyze_resume(
    resume: UploadFile = File(..., description="The resume file (PDF or DOCX)."),
    job_description: str = Form(..., description="The job description text."),
    interview_level: str = Form(
        "General Match (Overall Fit)",
        description="The specific interview focus level."
    )
) -> JSONResponse:
    """
    Analyzes a resume against a job description and provides skill matches,
    missing skills, and proficiency levels.

    Args:
        resume (UploadFile): The candidate's resume in PDF or DOCX format.
        job_description (str): The text of the job description.
        interview_level (str): The selected interview focus level (e.g., "Junior", "Senior").

    Returns:
        JSONResponse: A JSON object containing skill analysis data.

    Raises:
        HTTPException: If the file type is unsupported, Ollama server is unreachable,
                       or LLM response cannot be parsed.
    """
    resume_text = "" # Initialize to ensure it's always a string
    try:
        raw_extracted_text = await extract_resume_text(resume)
        
        # --- FIX: Ensure resume_text is always a string before slicing/concatenation ---
        if isinstance(raw_extracted_text, str):
            resume_text = raw_extracted_text
        else:
            logger.error(f"Extracted resume text is not a string. Type: {type(raw_extracted_text)}, Value: {raw_extracted_text}")
            # Attempt to convert to string, especially if it's a list (unexpected)
            if isinstance(raw_extracted_text, list):
                resume_text = "\n".join(map(str, raw_extracted_text))
                logger.warning("Converted list of extracted text to string by joining elements.")
            else:
                resume_text = str(raw_extracted_text) # Fallback for other non-string types
                logger.warning(f"Converted non-string type ({type(raw_extracted_text)}) to string using str().")
        # --- END FIX ---

    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Failed to extract resume text in /analyze: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to extract resume text: {str(e)}")

    MAX_RESUME_CHARS = 3000
    # This line should now be safe as resume_text is guaranteed to be a string
    resume_summary = resume_text[:MAX_RESUME_CHARS] + "\n...[truncated]" if len(resume_text) > MAX_RESUME_CHARS else resume_text

    # Prompt for skill analysis with detailed JSON structure request
    prompt = f"""
You are a professional technical interviewer and AI career coach.

Your task is to evaluate the candidate's resume against the job description and respond strictly in JSON.

**Interview Context:** {interview_level}

**Resume:**
{resume_summary}

**Job Description:**
{job_description}

The response should contain the following structured JSON keys:

1. "skills" – list of skill evaluations with:
   - "skill" (string): The name of the skill.
   - "matched" (boolean): true if the skill is clearly present in the resume, false otherwise.
   - "proficiency" (string): Estimated proficiency (Beginner/Intermediate/Advanced/Expert).
   - "confidence_level" (string): Your assessment of the candidate's confidence in this skill based on resume (High/Moderate/Unknown).

2. "missing_skills" – (list of strings): Skills from the Job Description that are NOT clearly present in the resume.

3. "interview_points" – (list of objects): Behavioral/technical points to prepare with:
   - "topic" (string): The subject of the interview point (e.g., "Teamwork", "Problem Solving with Python").
   - "suggested_star_response" (object): A structured STAR response example:
       - "situation" (string)
       - "task" (string)
       - "action" (string)
       - "result" (string)

4. "project_talking_points" – (list of objects): For each major project in the resume, highlight:
   - "project_name" (string): Name of the project.
   - "highlight" (list of strings): 1-2 key bullet points to discuss about the project.

5. "recommended_learning" – (list of objects): For each missing skill, list:
   - "skill" (string): The missing skill.
   - "resource" (string): A suggested learning resource (e.g., "Coursera: Deep Learning Specialization").
   - "type" (string): Type of resource (Official Docs / Hands-on Tutorial / Course / Book).

6. "questions_to_ask" – (list of strings): 2–4 thoughtful questions the candidate can ask the interviewer.

Strictly return a valid JSON object only, no extra text, no markdown code blocks outside the JSON.
"""

    try:
        # Make request to Ollama
        ollama_response = requests.post(
            "http://localhost:11434/api/chat",
            json={
                "model": "openchat",  # Or your preferred model (e.g., "llama2", "mistral")
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {
                    "temperature": 0.3  # Lower temperature for more factual extraction
                }
            },
            timeout=540  # Increased timeout for potentially longer responses
        )
        ollama_response.raise_for_status()  # Raise HTTPError for bad responses (4xx or 5xx)

        response_data = ollama_response.json()
        content = response_data.get("message", {}).get("content", "").strip()

        # Clean content before JSON parsing (removes common LLM artifacts)
        cleaned_content = re.sub(
            r',?\s*The resource is not available, please try again later\.', 
            '', 
            content
        )

        try:
            # Attempt to parse the JSON content from LLM response
            parsed_json = json.loads(cleaned_content)
            return JSONResponse(content=parsed_json)
        except json.JSONDecodeError:
            logger.error(f"Failed to decode JSON from LLM response: {cleaned_content}", exc_info=True)
            return JSONResponse(
                content={
                    "error": "Failed to parse LLM response as JSON. "
                    "This might be due to LLM outputting non-JSON text or invalid JSON.",
                    "raw_response": content # Keep raw content for debugging
                },
                status_code=500
            )

    except requests.exceptions.ConnectionError:
        logger.error("Could not connect to Ollama server.", exc_info=True)
        raise HTTPException(
            status_code=503,
            detail="Could not connect to Ollama server at http://localhost:11434. Please ensure Ollama is running."
        )
    except requests.exceptions.HTTPError as e:
        logger.error(f"Ollama API HTTP error: {e.response.status_code} - {e.response.text}", exc_info=True)
        raise HTTPException(
            status_code=e.response.status_code,
            detail=f"Error from Ollama API: {e.response.text}"
        )
    except Exception as e:
        logger.error(f"An unexpected error occurred during LLM analysis: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred during analysis: {str(e)}"
        )


@app.post("/interview_prep")
async def interview_prep(
    resume_text: str = Form(..., description="The extracted text from the resume."),
    job_description: str = Form(..., description="The job description text."),
    interview_level: str = Form(
        "General Match (Overall Fit)",
        description="The specific interview focus level."
    )
) -> JSONResponse:
    """
    Generates comprehensive interview preparation tips based on resume, job description,
    and interview level.

    Args:
        resume_text (str): The extracted text content of the resume.
        job_description (str): The text of the job description.
        interview_level (str): The selected interview focus level.

    Returns:
        JSONResponse: A JSON object containing the Markdown formatted tips.

    Raises:
        HTTPException: If Ollama server is unreachable or LLM response cannot be parsed.
    """
    # Prompt for interview preparation tips
    prompt = f"""
You are an expert career advisor and technical interviewer. Based on the following resume and job description, provide concise and actionable tips on how the candidate should prepare for an interview, specifically focusing on the "{interview_level}" context.

Your advice should cover:
1.  **Preparation for Matched Skills:** For each skill or area where the resume shows experience, suggest how to explain their experience, what common questions to expect, and what aspects to highlight.
2.  **Addressing Missing Skills:** For skills mentioned in the job description but not clearly on the resume, provide strategies on how to gracefully acknowledge this, what to say, and how to frame their learning agility or related experiences. This should include:
    * How to express eagerness to learn and growth mindset.
    * How to highlight transferable skills from other areas.
    * How to demonstrate a structured approach to acquiring new knowledge or tackling unfamiliar problems.
    * What to do if asked a direct question about an unknown skill (e.g., admit not knowing but explain how they would research or approach solving a problem with it).
    * Emphasis on demonstrating problem-solving abilities over specific tool knowledge.
3.  **General Interview Strategies:** Include broader tips relevant to the interview level. For example:
    * Importance of asking thoughtful questions.
    * Demonstrating cultural fit and teamwork.
    * Communicating thought process during technical questions.
    * Follow-up etiquette.

**Resume:**
{resume_text}

**Job Description:**
{job_description}

Please structure your response in a clear, easy-to-read Markdown format using prominent headings and bullet points. Do NOT return JSON for this request.
"""

    try:
        # Make request to Ollama
        ollama_response = requests.post(
            "http://localhost:11434/api/chat",
            json={
                "model": "openchat",  # Or your preferred model
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {
                    "temperature": 0.5  # A bit higher temperature for more creative tips
                }
            },
            timeout=180  # Increased timeout for potentially longer responses
        )
        ollama_response.raise_for_status()

        response_data = ollama_response.json()
        content = response_data.get("message", {}).get("content", "").strip()
        return JSONResponse(content={"tips": content})

    except requests.exceptions.ConnectionError:
        logger.error("Could not connect to Ollama server for interview tips.", exc_info=True)
        raise HTTPException(status_code=503, detail="Could not connect to Ollama server at http://localhost:11434. Please ensure Ollama is running.")
    except requests.exceptions.HTTPError as e:
        logger.error(f"Ollama API HTTP error for interview tips: {e.response.status_code} - {e.response.text}", exc_info=True)
        raise HTTPException(status_code=e.response.status_code, detail=f"Ollama API error: {e.response.text}")
    except Exception as e:
        logger.error(f"An unexpected error occurred during LLM analysis for interview tips: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during LLM analysis for interview tips: {str(e)}")